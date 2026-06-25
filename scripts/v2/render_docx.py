"""
v2 Word (.docx) 报告渲染器

依赖 python-docx（已安装）。复用 render_html.py 的 compute_report_context() 获取数据，
用 python-docx 生成 Word 文档。通过 --format docx 启用。
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

from scripts.v2.render_html import compute_report_context

# 异常值标红色
_RED = RGBColor(0xC8, 0x33, 0x33)
_DARK = RGBColor(0x33, 0x33, 0x33)
_ACCENT = RGBColor(0x1A, 0x73, 0xE8)


def _set_cell_text(cell, text: str, bold: bool = False, color: RGBColor = _DARK, size: int = 10) -> None:
    """设置表格单元格文本和样式。"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加标题，统一蓝色。"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = _ACCENT


def _add_table_head(doc: Document, headers: List[str], col_widths: Optional[List[float]] = None) -> Any:
    """创建带表头的表格，返回表格对象。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table


def _add_table_row(table: Any, cells: List[str], bold: bool = False,
                   abnormal: Optional[bool] = None) -> None:
    """添加表格行。abnormal=True 时该行数值标红。"""
    row = table.add_row()
    for i, val in enumerate(cells):
        color = _RED if (abnormal and i == 1) else _DARK
        _set_cell_text(row.cells[i], str(val), bold=bold, color=color)


# ---------------------------------------------------------------------------
# 核心导出函数
# ---------------------------------------------------------------------------

def render_docx_report(
    profile: Dict[str, Any],
    groups: Dict[str, List[Dict[str, Any]]],
    output_dir: Path,
) -> Optional[Path]:
    """渲染 .docx 报告。返回输出路径，失败返回 None。"""
    if not _DOCX_AVAILABLE:
        logger.warning('python-docx 未安装，跳过 DOCX 渲染')
        return None

    ctx = compute_report_context(profile, groups)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    # ======================== 封面 ========================
    dem = ctx.get('demographics', {})
    h = doc.add_heading('患者病情档案', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'患者：{dem.get("name", "—")}    性别：{dem.get("gender", "—")}    年龄：{dem.get("age", "—")}')
    doc.add_paragraph(f'主要诊断：{dem.get("primary_diagnosis", "—")}')
    doc.add_paragraph(f'建档日期：{ctx.get("updated_at", "")[:10]}')
    doc.add_paragraph('')

    # ======================== 病情速览 ========================
    if ctx.get('clinical_summary'):
        _add_heading(doc, '病情速览', level=1)
        doc.add_paragraph(ctx['clinical_summary'])

    # 关键指标快照
    critical = ctx.get('critical_alerts', [])
    if critical:
        _add_heading(doc, '⚠️ 危急值/异常警报', level=1)
        for c in critical:
            p = doc.add_paragraph()
            run = p.add_run(f'{c.get("emoji", "")} {c.get("item_name", "")}：{c.get("message", "")}')
            run.bold = True
            run.font.color.rgb = _RED

    # ======================== 诊疗时间线 ========================
    timeline = ctx.get('timeline', [])
    if timeline:
        _add_heading(doc, '诊疗时间线', level=1)
        t = _add_table_head(doc, ['日期', '事件', '关键结果/疗效'])
        for item in timeline:
            _add_table_row(t, [
                str(item.get('dates', '') or item.get('date', '')),
                str(item.get('title', '') or item.get('event', '')),
                str(item.get('note', '') or item.get('result', '')),
            ])

    tl_items = ctx.get('timeline_items', [])
    if tl_items:
        _add_heading(doc, '详细时间线', level=1)
        t = _add_table_head(doc, ['日期', '事件', '医院', '结果'])
        for item in tl_items:
            _add_table_row(t, [
                str(item.get('date', '')),
                str(item.get('event', '')),
                str(item.get('hospital', '')),
                str(item.get('result', '')),
            ])

    # ======================== 检查指标趋势 ========================
    tumor_marker_tables = ctx.get('tumor_marker_tables', {})
    if tumor_marker_tables:
        _add_heading(doc, '肿瘤标志物趋势', level=1)
        for marker_name, mdata in tumor_marker_tables.items():
            p = doc.add_paragraph()
            run = p.add_run(f'{marker_name}（{mdata.get("unit", "")}）')
            run.bold = True
            rows = mdata.get('rows', [])
            if rows:
                t = _add_table_head(doc, ['日期', f'{marker_name}({mdata.get("unit","")})'])
                for r in rows:
                    _add_table_row(t, [
                        str(r.get('date', '') or r.get('x', '')),
                        str(r.get('value', '') or r.get('y', '')),
                    ], abnormal=r.get('abnormal'))

    # 分类检验指标表
    cat_lab = ctx.get('categorized_lab', [])
    if cat_lab:
        _add_heading(doc, '检验指标', level=1)
        for cat in cat_lab:
            p = doc.add_paragraph()
            run = p.add_run(f'▸ {cat["category"]}（{len(cat["items"])}项）')
            run.bold = True
            t = _add_table_head(doc, ['指标', '数值', '单位', '参考范围'])
            for it in cat['items']:
                ref = ''
                if it.get('ref_low') is not None and it.get('ref_high') is not None:
                    ref = f"{it['ref_low']}–{it['ref_high']}"
                _add_table_row(t, [
                    str(it.get('name', '')),
                    str(it.get('value', '')),
                    str(it.get('unit', '')),
                    ref,
                ], abnormal=it.get('abnormal'))

    # ======================== 病理与基因 ========================
    pathology_diag = ctx.get('pathology_diagnosis')
    if pathology_diag:
        _add_heading(doc, '病理诊断', level=1)
        fields = [
            ('肿瘤部位', 'tumor_site'), ('肿瘤大小', 'tumor_size'),
            ('TNM分期', 'tnm_stage'), ('病理诊断', 'pathology_diagnosis'),
        ]
        for label, key in fields:
            if pathology_diag.get(key):
                p = doc.add_paragraph()
                run = p.add_run(f'{label}：')
                run.bold = True
                p.add_run(str(pathology_diag[key]))

    # 免疫组化
    ihc = ctx.get('ihc_items', [])
    if ihc:
        _add_heading(doc, '免疫组化', level=2)
        t = _add_table_head(doc, ['指标', '结果', '临床意义'])
        for m in ihc:
            _add_table_row(t, [
                str(m.get('marker', '') or m.get('name', '')),
                str(m.get('result', '')),
                str(m.get('clinical_meaning', '') or m.get('meaning', '')),
            ])

    # 基因检测
    genetic = ctx.get('genetic_highlights', [])
    if genetic:
        _add_heading(doc, '基因检测', level=2)
        t = _add_table_head(doc, ['基因', '突变', '类型', '致病性', '证据等级'])
        for g in genetic:
            _add_table_row(t, [
                str(g.get('gene', '') or g.get('marker', '')),
                str(g.get('mutation', '') or g.get('result', '')),
                str(g.get('category', '') or g.get('type', '')),
                str('致病' if g.get('pathogenic') else '意义未明'),
                str(g.get('evidence_tier', '') or g.get('tier_label', '')),
            ])

    # ======================== 用药方案 ========================
    med_table = ctx.get('medication_table', [])
    if med_table:
        _add_heading(doc, '当前用药', level=1)
        t = _add_table_head(doc, ['药物', '剂量', '给药方式', '用途'])
        for m in med_table:
            _add_table_row(t, [
                str(m.get('name', '')),
                str(m.get('dose', '') or m.get('dosage', '')),
                str(m.get('route', '')),
                str(m.get('purpose', '')),
            ])

    med = ctx.get('medication', {})
    if med:
        _add_heading(doc, '用药方案汇总', level=2)
        current = med.get('current', [])
        history = med.get('history', [])
        if current:
            p = doc.add_paragraph()
            run = p.add_run('当前用药：')
            run.bold = True
            p.add_run('、'.join(current))
        if history:
            p = doc.add_paragraph()
            run = p.add_run('历史用药：')
            run.bold = True
            p.add_run('、'.join(history))

    # ======================== 影像检查 ========================
    imaging = ctx.get('imaging_summary', [])
    if imaging:
        _add_heading(doc, '影像检查', level=1)
        t = _add_table_head(doc, ['日期', '检查项目', '关键发现（对比前片）'])
        for im in imaging:
            _add_table_row(t, [
                str(im.get('date', '')),
                str(im.get('modality', '')),
                str(im.get('findings', '')),
            ])

    # ======================== 关注问题要点 ========================
    concerns = ctx.get('key_concerns', [])
    if concerns:
        _add_heading(doc, '关注问题要点', level=1)
        for c in concerns:
            if isinstance(c, dict):
                p = doc.add_paragraph()
                run = p.add_run(f'• {c.get("concern", "") or c.get("item_name", "")}：')
                run.bold = True
                p.add_run(str(c.get('detail', '') or c.get('message', '')))
            else:
                doc.add_paragraph(f'• {c}')

    # ======================== 问诊咨询建议 ========================
    questions = ctx.get('consultation_questions', [])
    if questions:
        _add_heading(doc, '问诊咨询建议', level=1)
        for q in questions:
            doc.add_paragraph(f'{q}', style='List Bullet')

    # ======================== 信息缺口提示 ========================
    gaps = ctx.get('gaps', [])
    if gaps:
        _add_heading(doc, '信息缺口提示', level=1)
        for g in gaps:
            doc.add_paragraph(f'⚠️ {g}')

    # ======================== 附件目录 ========================
    files = ctx.get('files', [])
    if files:
        _add_heading(doc, '附件目录', level=1)
        t = _add_table_head(doc, ['文件名', '日期', '分类'])
        for f in files:
            _add_table_row(t, [
                str(f.get('title', '')),
                str(f.get('date', '')),
                str(f.get('category', '')),
            ])

    # ======================== 免责声明 ========================
    doc.add_paragraph('')
    h = doc.add_heading('免责声明', level=2)
    for run in h.runs:
        run.font.color.rgb = _RED
    doc.add_paragraph(
        '本报告由 AI 辅助整理，仅供医疗专业人士参考。不构成医学诊断或治疗建议。'
        '所有检验结果以原始报告为准，异常值请结合临床判断。'
        '如遇危急值，请立即联系主管医生或急诊处理。'
    )

    # 保存
    output_path = Path(output_dir) / 'report.docx'
    doc.save(str(output_path))
    logger.info('DOCX 报告已生成: %s', output_path)
    return output_path
